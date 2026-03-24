#!/usr/bin/env python3
"""Generate Clinvara User Guide as a Word document."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

# --- Styles ---
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)

for level in range(1, 4):
    h = doc.styles[f"Heading {level}"]
    h.font.name = "Calibri"
    h.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)  # slate-900


def add_tip(text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(f"Tip: {text}")
    run.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x47, 0x55, 0x69)


def add_steps(steps: list[str]) -> None:
    for i, step in enumerate(steps, 1):
        doc.add_paragraph(f"{i}. {step}", style="List Number")


def add_bullet_list(items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


# ============================================================
# COVER PAGE
# ============================================================
for _ in range(6):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("Clinvara")
run.font.size = Pt(36)
run.bold = True
run.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run("User Guide")
run.font.size = Pt(20)
run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

version = doc.add_paragraph()
version.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = version.add_run("Version 1.0 — March 2026")
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)

doc.add_page_break()

# ============================================================
# TABLE OF CONTENTS (manual)
# ============================================================
doc.add_heading("Contents", level=1)
toc_items = [
    "1. Getting Started",
    "2. Dashboard",
    "3. Managing Patients",
    "4. Care Plans",
    "5. Assessments",
    "6. Medications",
    "7. Patient Flow",
    "8. Workforce Management",
    "9. Practitioners",
    "10. Billing & Subscription",
    "11. Settings",
    "12. Team Management",
    "13. Tenant Management (Super Admin)",
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(4)

doc.add_page_break()

# ============================================================
# 1. GETTING STARTED
# ============================================================
doc.add_heading("1. Getting Started", level=1)

doc.add_heading("Logging In", level=2)
doc.add_paragraph(
    "To access Clinvara, open your web browser and navigate to your organisation's Clinvara URL."
)
add_steps([
    "Enter your email address and password.",
    "Click the Log In button.",
    "If this is your first login, you will be asked to change your password.",
    "You will be taken to the Dashboard.",
])
add_tip("If you belong to more than one organisation, you will be asked to select which one to work in.")

doc.add_heading("Navigating the Application", level=2)
doc.add_paragraph(
    "The left sidebar contains links to all the main sections of the application. "
    "Click any item to navigate to that section. The sidebar shows your name and role at the bottom."
)
add_bullet_list([
    "Dashboard — Overview of key statistics and recent activity",
    "Patients — View and manage patient records",
    "Care Plans — Create and track care plans",
    "Medications — Manage prescriptions",
    "Assessments — Record clinical assessments",
    "Roster — View and manage staff schedules",
    "Patient Flow — Monitor bed occupancy and encounters",
    "Locations & Beds — Manage wards and beds",
    "Practitioners — Manage clinical staff records",
])

doc.add_heading("Changing Your Password", level=2)
doc.add_paragraph(
    'Click "Change Password" in the sidebar to update your password at any time.'
)

doc.add_page_break()

# ============================================================
# 2. DASHBOARD
# ============================================================
doc.add_heading("2. Dashboard", level=1)
doc.add_paragraph(
    "The Dashboard gives you a quick overview of your organisation's activity."
)
doc.add_heading("What You Will See", level=2)
add_bullet_list([
    "Active Patients — Total number of active patients in your organisation",
    "Practitioners — Number of registered practitioners",
    "Users — Number of staff users on the platform",
    "Shifts This Week — Number of scheduled shifts for the current week",
    "Active Encounters — Patients currently admitted",
    "Available Beds — Beds ready for use",
    "Recent Patients — Quick links to recently added patients",
    "Patient Demographics — Breakdown of patients by gender",
    "Recent Activity — Timeline of recent events (admissions, notes, assessments, etc.)",
])
add_tip("Click on a patient name to go directly to their record.")

doc.add_page_break()

# ============================================================
# 3. MANAGING PATIENTS
# ============================================================
doc.add_heading("3. Managing Patients", level=1)

doc.add_heading("Viewing Patients", level=2)
doc.add_paragraph(
    'Go to "Patients" in the sidebar. You will see a table of all patients with their name, '
    "NHS Number, date of birth, gender, and organisation."
)
add_bullet_list([
    "Use the search box to find patients by name or NHS Number.",
    "Click a patient's name to view their full record.",
    "Use the Previous/Next buttons at the bottom to move between pages.",
])

doc.add_heading("Adding a New Patient", level=2)
doc.add_paragraph(
    'Click the "New Patient" button in the top-right corner.'
)
add_steps([
    "Fill in the required fields marked with an asterisk (*):",
    "   - First Name and Last Name",
    "   - Gender and Date of Birth",
    "Optionally add contact details (phone, email), address, and identifiers (NHS Number, MRN).",
    'Click "Register Patient" to save.',
])
add_tip("NHS Numbers must be exactly 10 digits. UK phone numbers and postcodes are validated automatically.")

doc.add_heading("Viewing a Patient Record", level=2)
doc.add_paragraph(
    "A patient's record shows all their information in one place:"
)
add_bullet_list([
    "Personal Information — Name, date of birth, gender, contact details",
    "Address — Full postal address",
    "Care Team — Assigned GP/practitioner and managing organisation",
    "Care Plans — All care plans for this patient",
    "Assessments — Clinical assessments with risk levels",
    "Medications — Active prescriptions",
    "Timeline — A chronological record of all events",
])

doc.add_heading("Editing Patient Details", level=2)
add_steps([
    'Click the "Edit Patient" button at the top of the patient record.',
    "Update the fields you need to change.",
    'Click "Save Changes" to save.',
])

doc.add_heading("Adding a Timeline Event", level=2)
add_steps([
    'On the patient record, click "Add Event" in the Timeline section.',
    "Select the event type (Note, Admission, Discharge, Transfer, Assessment, or Referral).",
    "Enter a summary of the event.",
    'Click "Add Event" to save.',
])

doc.add_page_break()

# ============================================================
# 4. CARE PLANS
# ============================================================
doc.add_heading("4. Care Plans", level=1)

doc.add_heading("Viewing Care Plans", level=2)
doc.add_paragraph(
    'Go to "Care Plans" in the sidebar. The list shows all care plans with their title, patient, '
    "category, status, start date, and author."
)
add_bullet_list([
    "Filter by Status: Draft, Active, Completed, or Cancelled",
    "Filter by Category: General, Nursing, Physiotherapy, Mental Health, or Palliative",
    "Click a care plan title to view its details.",
])

doc.add_heading("Creating a Care Plan", level=2)
add_steps([
    'Click "New Care Plan".',
    "Select the patient this care plan is for.",
    "Enter a title and description.",
    "Choose a category.",
    "Set the start date (required) and optionally an end date and review date.",
    'Click "Create Care Plan".',
])
add_tip('You can also create a care plan directly from a patient\'s record by clicking "+ New" in their Care Plans section.')

doc.add_heading("Care Plan Statuses", level=2)
table = doc.add_table(rows=5, cols=2)
table.style = "Light Grid Accent 1"
table.rows[0].cells[0].text = "Status"
table.rows[0].cells[1].text = "Meaning"
statuses = [
    ("Draft", "Plan is being prepared, not yet active"),
    ("Active", "Plan is currently in use"),
    ("Completed", "Plan has been completed successfully"),
    ("Cancelled", "Plan was cancelled"),
]
for i, (status, meaning) in enumerate(statuses, 1):
    table.rows[i].cells[0].text = status
    table.rows[i].cells[1].text = meaning

doc.add_page_break()

# ============================================================
# 5. ASSESSMENTS
# ============================================================
doc.add_heading("5. Assessments", level=1)

doc.add_heading("Viewing Assessments", level=2)
doc.add_paragraph(
    "The Assessments page shows all clinical assessments. Each row displays the title, patient, "
    "type, score, risk level, status, and date."
)
add_bullet_list([
    "Filter by assessment type (e.g. Waterlow, MUST, Tinetti)",
    "Filter by risk level: No Risk, Low, Medium, High, Very High",
])

doc.add_heading("Creating an Assessment", level=2)
add_steps([
    'Click "New Assessment".',
    "Select the patient.",
    "Enter a title and choose the assessment type.",
    "Optionally specify the tool name (e.g. Waterlow Scale).",
    "Enter the score, maximum score, and risk level.",
    "Add clinical notes and recommended actions.",
    'Click "Create Assessment".',
])

doc.add_heading("Risk Levels", level=2)
table = doc.add_table(rows=6, cols=2)
table.style = "Light Grid Accent 1"
table.rows[0].cells[0].text = "Risk Level"
table.rows[0].cells[1].text = "Description"
risks = [
    ("No Risk", "No risk factors identified"),
    ("Low", "Minor risk — routine monitoring"),
    ("Medium", "Moderate risk — care plan may be needed"),
    ("High", "Significant risk — intervention required"),
    ("Very High", "Critical risk — immediate action needed"),
]
for i, (level, desc) in enumerate(risks, 1):
    table.rows[i].cells[0].text = level
    table.rows[i].cells[1].text = desc

doc.add_page_break()

# ============================================================
# 6. MEDICATIONS
# ============================================================
doc.add_heading("6. Medications", level=1)

doc.add_heading("Viewing Prescriptions", level=2)
doc.add_paragraph(
    "The Medications page lists all prescriptions. Each row shows the medication name, patient, "
    "dosage, status, start date, and prescriber."
)
doc.add_paragraph("Use the status filter to narrow results: Draft, Active, On Hold, Completed, Stopped, or Cancelled.")

doc.add_heading("Creating a Prescription", level=2)
add_steps([
    'Click "New Prescription".',
    "Select the patient.",
    "Choose or enter the medication name.",
    "Enter the dosage instructions and frequency.",
    "Set start and end dates.",
    "Select the prescriber.",
    'Click "Create Prescription".',
])
add_tip('You can also create a prescription from a patient\'s record by clicking "+ New" in their Medications section.')

doc.add_heading("Prescription Statuses", level=2)
table = doc.add_table(rows=7, cols=2)
table.style = "Light Grid Accent 1"
table.rows[0].cells[0].text = "Status"
table.rows[0].cells[1].text = "Meaning"
med_statuses = [
    ("Draft", "Prescription is being prepared"),
    ("Active", "Medication is currently prescribed"),
    ("On Hold", "Temporarily paused"),
    ("Completed", "Course of medication finished"),
    ("Stopped", "Medication discontinued"),
    ("Cancelled", "Prescription was cancelled"),
]
for i, (status, meaning) in enumerate(med_statuses, 1):
    table.rows[i].cells[0].text = status
    table.rows[i].cells[1].text = meaning

doc.add_page_break()

# ============================================================
# 7. PATIENT FLOW
# ============================================================
doc.add_heading("7. Patient Flow", level=1)
doc.add_paragraph(
    "Patient Flow helps you manage bed occupancy, patient admissions, transfers, and discharges."
)

doc.add_heading("Patient Flow Dashboard", level=2)
doc.add_paragraph("The dashboard shows at a glance:")
add_bullet_list([
    "Active Encounters — Number of currently admitted patients",
    "Available Beds — Beds ready for new admissions",
    "Occupied Beds — Beds currently in use",
    "Occupancy Rate — Percentage of beds occupied",
    "Bed Overview — Visual map of beds by location, colour-coded by status",
    "Active Encounters list — Quick access to current patient encounters",
])

doc.add_heading("Bed Status Colours", level=2)
table = doc.add_table(rows=5, cols=2)
table.style = "Light Grid Accent 1"
table.rows[0].cells[0].text = "Colour"
table.rows[0].cells[1].text = "Status"
bed_colours = [
    ("Green", "Available"),
    ("Red", "Occupied"),
    ("Yellow", "Maintenance"),
    ("Grey", "Closed"),
]
for i, (colour, status) in enumerate(bed_colours, 1):
    table.rows[i].cells[0].text = colour
    table.rows[i].cells[1].text = status

doc.add_heading("Admitting a Patient", level=2)
add_steps([
    'Click "Admit Patient" on the Patient Flow dashboard.',
    "Select the patient to admit.",
    "Choose the encounter class (Inpatient, Outpatient, Emergency, or Home Health).",
    "Optionally select the admission source.",
    "Choose a location and bed.",
    "Add any notes.",
    'Click "Admit Patient".',
])

doc.add_heading("Transferring a Patient", level=2)
add_steps([
    "Open the encounter from the Patient Flow dashboard.",
    'Click "Transfer Patient".',
    "Select the destination location and bed.",
    "Enter a reason for the transfer.",
    'Click "Confirm Transfer".',
])

doc.add_heading("Discharging a Patient", level=2)
add_steps([
    "Open the encounter from the Patient Flow dashboard.",
    'Click "Discharge Patient".',
    "Select the discharge destination (Home, Care Home, Hospital Transfer, etc.).",
    "Add any discharge notes.",
    'Click "Confirm Discharge".',
])

doc.add_heading("Managing Locations & Beds", level=2)
doc.add_paragraph(
    'Go to "Locations & Beds" in the sidebar.'
)
add_bullet_list([
    'Click "Add Location" to create a new ward, room, or department.',
    "Enter the location name, type, ward, floor, and capacity.",
    'Click "Add Bed" on a location card to add beds to that location.',
    "Beds are shown as coloured squares indicating their status.",
])

doc.add_page_break()

# ============================================================
# 8. WORKFORCE MANAGEMENT
# ============================================================
doc.add_heading("8. Workforce Management", level=1)

doc.add_heading("Roster", level=2)
doc.add_paragraph(
    "The Roster page shows a weekly calendar view of all scheduled shifts."
)
add_bullet_list([
    "Each column represents a day of the week.",
    "Shifts are colour-coded by type (Early, Late, Night, Long Day, Twilight).",
    "Filter by location to see shifts for a specific ward.",
    "Unassigned shifts are shown separately.",
])

doc.add_heading("Shift Types", level=2)
table = doc.add_table(rows=7, cols=2)
table.style = "Light Grid Accent 1"
table.rows[0].cells[0].text = "Type"
table.rows[0].cells[1].text = "Description"
shift_types = [
    ("Early", "Morning shift"),
    ("Late", "Afternoon/evening shift"),
    ("Night", "Overnight shift"),
    ("Long Day", "Extended day shift"),
    ("Twilight", "Late evening shift"),
    ("Custom", "Custom shift pattern"),
]
for i, (stype, desc) in enumerate(shift_types, 1):
    table.rows[i].cells[0].text = stype
    table.rows[i].cells[1].text = desc

doc.add_heading("Shift Patterns (Admin)", level=2)
doc.add_paragraph(
    "Shift Patterns are reusable templates that define standard shift times."
)
add_steps([
    'Go to "Shift Patterns" in the sidebar.',
    'Click "New Pattern".',
    "Enter a name, select the shift type, and set start/end times and break duration.",
    'Click "Create Pattern".',
])

doc.add_heading("Availability", level=2)
doc.add_paragraph(
    "Staff members can mark their availability using the Availability page. "
    "This helps managers plan the roster around staff schedules."
)

doc.add_heading("Shift Swaps", level=2)
doc.add_paragraph(
    "The Shift Swap Marketplace allows staff to request and manage shift trades."
)
add_bullet_list([
    "Request a swap — Offer one of your assigned shifts for trade.",
    "Browse available swaps — See shifts offered by colleagues.",
    "Accept a swap — Offer to take a colleague's shift.",
    "Manager approval — An admin must approve the swap before it takes effect.",
])

doc.add_page_break()

# ============================================================
# 9. PRACTITIONERS
# ============================================================
doc.add_heading("9. Practitioners", level=1)
doc.add_paragraph(
    "The Practitioners page lets you manage clinical staff records."
)

doc.add_heading("Adding a Practitioner", level=2)
add_steps([
    'Click "Add Practitioner".',
    "Enter the practitioner's first name, last name, and gender.",
    "Add contact details (phone, email).",
    "Select their specialty.",
    "Enter their registration number (e.g. GMC, NMC number).",
    'Click "Save".',
])

doc.add_heading("Editing a Practitioner", level=2)
doc.add_paragraph(
    'Click the "Edit" button next to a practitioner to update their details.'
)

doc.add_heading("Deactivating a Practitioner", level=2)
doc.add_paragraph(
    'Click "Deactivate" next to a practitioner to mark them as inactive. '
    "They will remain in the system for record-keeping but will not appear in active lists."
)

doc.add_page_break()

# ============================================================
# 10. BILLING
# ============================================================
doc.add_heading("10. Billing & Subscription", level=1)
doc.add_paragraph(
    'Go to "Billing" in the sidebar to view your subscription and manage your plan.'
)

doc.add_heading("Your Current Plan", level=2)
doc.add_paragraph(
    "The top of the page shows your current subscription, including:"
)
add_bullet_list([
    "Plan name and status",
    "Patient and user limits",
    "Renewal or cancellation date",
])

doc.add_heading("Available Plans", level=2)
table = doc.add_table(rows=5, cols=4)
table.style = "Light Grid Accent 1"
headers = ["Plan", "Price", "Patients", "Users"]
for i, h in enumerate(headers):
    table.rows[0].cells[i].text = h
plans = [
    ("Free", "Free", "5", "1"),
    ("Starter", "£59/month", "200", "20"),
    ("Professional", "£99/month", "500", "50"),
    ("Enterprise", "£299/month", "Unlimited", "Unlimited"),
]
for i, (plan, price, patients, users) in enumerate(plans, 1):
    table.rows[i].cells[0].text = plan
    table.rows[i].cells[1].text = price
    table.rows[i].cells[2].text = patients
    table.rows[i].cells[3].text = users

doc.add_paragraph()

doc.add_heading("Upgrading Your Plan", level=2)
doc.add_paragraph(
    'Click "Upgrade" on the plan you want. You will be redirected to a secure payment page '
    "powered by Stripe to complete the upgrade."
)

doc.add_heading("Managing Billing", level=2)
doc.add_paragraph(
    'Click "Manage Billing" to access the Stripe billing portal where you can update your '
    "payment method, view invoices, and cancel your subscription."
)

doc.add_page_break()

# ============================================================
# 11. SETTINGS
# ============================================================
doc.add_heading("11. Settings (Admin Only)", level=1)
doc.add_paragraph(
    "The Settings page is available to administrators and provides access to configuration options."
)

doc.add_heading("Organisation Settings", level=2)
doc.add_paragraph(
    "Update your organisation's name, address, contact details, and type."
)

doc.add_heading("Assessment Types", level=2)
doc.add_paragraph(
    "Configure the types of clinical assessments available in your organisation "
    "(e.g. Waterlow, MUST, Tinetti, Falls Risk). These appear as options when creating assessments."
)

doc.add_heading("Specialty Types", level=2)
doc.add_paragraph(
    "Manage the list of practitioner specialties available when adding or editing practitioners."
)

doc.add_heading("Medication Types", level=2)
doc.add_paragraph(
    "Manage your medication catalogue. Add medications with their BNF code, strength, and form. "
    "These appear as options when creating prescriptions."
)

doc.add_page_break()

# ============================================================
# 12. TEAM MANAGEMENT
# ============================================================
doc.add_heading("12. Team Management (Admin Only)", level=1)
doc.add_paragraph(
    'Go to "Team" in the sidebar to manage users in your organisation.'
)
add_bullet_list([
    "View all team members and their roles.",
    "Invite new users to join your organisation.",
    "Assign roles (Standard User or Administrator).",
    "Remove team members.",
])
add_tip("User limits are based on your subscription plan. Upgrade your plan if you need more users.")

doc.add_page_break()

# ============================================================
# 13. TENANT MANAGEMENT
# ============================================================
doc.add_heading("13. Tenant Management (Super Admin Only)", level=1)
doc.add_paragraph(
    "This section is only available to Super Administrators who manage the entire Clinvara platform."
)

doc.add_heading("Managing Tenants", level=2)
doc.add_paragraph(
    "The Tenants page shows all organisations on the platform. Each card displays the "
    "organisation name, type, status, subscription tier, and contact details."
)
doc.add_paragraph("Click on a tenant card to select it and work within that organisation's context.")

doc.add_heading("Managing Admins", level=2)
add_bullet_list([
    "Super Admins — Create and manage platform-wide administrators.",
    "Tenant Admins — Create and manage administrators for specific organisations.",
])

# ============================================================
# FOOTER
# ============================================================
doc.add_page_break()
footer = doc.add_paragraph()
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = footer.add_run("Clinvara User Guide v1.0")
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)
p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = p2.add_run("© 2026 Clinvara. All rights reserved.")
run2.font.size = Pt(10)
run2.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)

# Save
output_path = "/Users/Temmi/Documents/Development/saas-products/care-solutions/docs/Clinvara-User-Guide.docx"
import os
os.makedirs(os.path.dirname(output_path), exist_ok=True)
doc.save(output_path)
print(f"Saved to {output_path}")
